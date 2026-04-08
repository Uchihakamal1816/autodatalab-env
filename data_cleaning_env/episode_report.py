# Copyright (c) Meta Platforms, Inc. and affiliates.
# All rights reserved.
#
# This source code is licensed under the BSD-style license found in the
# LICENSE file in the root directory of this source tree.

"""Word (.docx) reports: final outcome per task (operations + resulting table)."""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Optional

from data_cleaning_env.models import DataCleaningAction, DataCleaningObservation


@dataclass
class EpisodeTrace:
    """Final-outcome summary for one task (Word export)."""

    task: str
    mode: str
    model_name: Optional[str]
    instruction: str
    terminal_grader_score: Optional[float]
    operations_done: List[str] = field(default_factory=list)
    final_row_count: Optional[int] = None
    final_columns: List[str] = field(default_factory=list)
    final_preview: List[Dict[str, Any]] = field(default_factory=list)
    remaining_issues: List[str] = field(default_factory=list)
    remaining_policy_warnings: List[str] = field(default_factory=list)


@dataclass
class _StepSnap:
    action: Dict[str, Any]


def _brief_action_dict(d: Dict[str, Any]) -> str:
    at = d.get("action_type", "?")
    if at == "remove_duplicates":
        return "Remove duplicate rows"
    if at == "fill_missing":
        col = d.get("column") or "?"
        m = d.get("method") or "?"
        return f"Fill missing values in {col} ({m})"
    if at == "drop_column":
        return f"Drop column {d.get('column', '?')}"
    if at == "normalize":
        return f"Normalize column {d.get('column', '?')}"
    if at == "remove_outliers":
        z = d.get("z_threshold", 3.0)
        return f"Remove outliers in {d.get('column', '?')} (z<={z})"
    if at == "plot":
        return f"Declare plot: {d.get('plot_type')} {d.get('x')} vs {d.get('y')}"
    if at == "derive_revenue":
        return "Derive Revenue = Price x Quantity"
    if at == "compute_metrics":
        return "Aggregate revenue by Category (e-commerce metrics)"
    if at == "compute_kpis":
        return "Compute business KPIs: TotalRevenue + AvgOrderValue"
    if at == "export_csv":
        return "Export working table to CSV"
    if at == "submit":
        return "Submit (finish episode)"
    if at == "noop":
        return "No-op"
    return str(at)


def _ensure_docx() -> None:
    try:
        import docx  # noqa: F401
    except ImportError as e:  # pragma: no cover
        raise ImportError(
            "Word reports require: pip install python-docx (or pip install -e \".[report]\")"
        ) from e


def _set_section_margins(doc: Any) -> None:
    from docx.shared import Inches

    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)


def _shade_table_header(table: Any, fill: str = "D9E2F3") -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)


def _add_preview_table(doc: Any, records: List[Dict[str, Any]], max_rows: int = 15) -> None:
    from docx.shared import Pt

    if not records:
        p = doc.add_paragraph("(No rows to preview.)")
        p.paragraph_format.space_after = Pt(12)
        return
    cols = list(records[0].keys())
    n_cols = len(cols)
    n_body = min(len(records), max_rows)
    table = doc.add_table(rows=1 + n_body, cols=n_cols)
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, col_name in enumerate(cols):
        hdr[i].text = str(col_name)
        for para in hdr[i].paragraphs:
            for r in para.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(records[:max_rows], start=1):
        for ci, c in enumerate(cols):
            v = row.get(c, "")
            table.rows[ri].cells[ci].text = "" if v is None else str(v)
            for para in table.rows[ri].cells[ci].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)
    _shade_table_header(table)
    doc.add_paragraph()


def _add_bullet_list(doc: Any, items: List[str]) -> None:
    from docx.shared import Pt

    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.size = Pt(11)


def _build_single_task_document(trace: EpisodeTrace) -> Any:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()
    _set_section_margins(doc)
    doc.core_properties.title = "Data cleaning: final report"
    doc.core_properties.subject = f"task={trace.task}"
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    h0 = doc.add_heading("Data cleaning: final report", level=0)
    h0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h0.runs:
        run.font.size = Pt(20)
        run.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(datetime.now(timezone.utc).strftime("%d %B %Y, %H:%M UTC"))
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

    doc.add_paragraph()

    score_txt = (
        f"{trace.terminal_grader_score:.4f}"
        if trace.terminal_grader_score is not None
        else "n/a"
    )
    meta = doc.add_paragraph()
    meta.add_run("Task: ").bold = True
    meta.add_run(f"{trace.task}    ")
    meta.add_run("Mode: ").bold = True
    meta.add_run(f"{trace.mode}    ")
    meta.add_run("Grader score: ").bold = True
    meta.add_run(score_txt)
    if trace.model_name:
        meta.add_run("    ")
        meta.add_run("Model: ").bold = True
        meta.add_run(trace.model_name)

    doc.add_heading("Benchmark instruction", level=1)
    p_inst = doc.add_paragraph(trace.instruction)
    p_inst.paragraph_format.space_after = Pt(12)
    for run in p_inst.runs:
        run.font.size = Pt(11)

    doc.add_heading("Operations performed", level=1)
    _add_bullet_list(doc, trace.operations_done or ["(none)"])

    doc.add_heading("Final dataset", level=1)
    p_sum = doc.add_paragraph()
    if trace.final_row_count is not None:
        p_sum.add_run("Rows: ").bold = True
        p_sum.add_run(str(trace.final_row_count))
        p_sum.add_run("    ")
    if trace.final_columns:
        p_sum.add_run("Columns: ").bold = True
        p_sum.add_run(", ".join(trace.final_columns))
    if not trace.final_columns and trace.final_row_count is None:
        p_sum.add_run("(Snapshot not available.)")

    doc.add_heading("Preview (first rows)", level=1)
    _add_preview_table(doc, trace.final_preview)

    doc.add_heading("Checks after run", level=1)
    issues = ", ".join(trace.remaining_issues) if trace.remaining_issues else "None"
    pol = (
        "; ".join(trace.remaining_policy_warnings)
        if trace.remaining_policy_warnings
        else "None"
    )
    p_chk = doc.add_paragraph()
    p_chk.add_run("Issue flags: ").bold = True
    p_chk.add_run(issues)
    p_chk.add_run("\n")
    p_chk.add_run("Policy warnings: ").bold = True
    p_chk.add_run(pol)

    return doc


def write_episode_docx(trace: EpisodeTrace, path: Path | str) -> Path:
    """Write one-task final report as Word. Requires ``python-docx``."""
    _ensure_docx()
    path = Path(path).expanduser().resolve()
    if path.suffix.lower() != ".docx":
        path = path.with_suffix(".docx")
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = _build_single_task_document(trace)
    doc.save(str(path))
    return path


def write_session_docx(traces: List[EpisodeTrace], path: Path | str) -> Path:
    """Combined report: cover page + one section per task."""
    _ensure_docx()
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    path = Path(path).expanduser().resolve()
    if path.suffix.lower() != ".docx":
        path = path.with_suffix(".docx")
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_section_margins(doc)
    doc.core_properties.title = "AutoDataLab: session report"
    doc.core_properties.subject = "Multi-task cleaning summary"
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_heading("AutoDataLab: session report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)
        run.font.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = sub.add_run(datetime.now(timezone.utc).strftime("%d %B %Y, %H:%M UTC"))
    rr.italic = True
    rr.font.size = Pt(10)
    rr.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

    task_list = doc.add_paragraph()
    task_list.alignment = WD_ALIGN_PARAGRAPH.CENTER
    task_list.add_run("Tasks covered: ").bold = True
    task_list.add_run(", ".join(t.task for t in traces))

    for ti, trace in enumerate(traces):
        if ti > 0:
            doc.add_page_break()
        doc.add_heading(f"Task: {trace.task}", level=1)
        sc = (
            f"{trace.terminal_grader_score:.4f}"
            if trace.terminal_grader_score is not None
            else "n/a"
        )
        m = doc.add_paragraph()
        m.add_run("Mode: ").bold = True
        m.add_run(f"{trace.mode}    ")
        m.add_run("Grader: ").bold = True
        m.add_run(sc)
        if trace.model_name:
            m.add_run("    ")
            m.add_run("Model: ").bold = True
            m.add_run(trace.model_name)

        doc.add_heading("Instruction", level=2)
        doc.add_paragraph(trace.instruction)

        doc.add_heading("Operations", level=2)
        _add_bullet_list(doc, trace.operations_done or ["(none)"])

        doc.add_heading("Result", level=2)
        rtxt = doc.add_paragraph()
        if trace.final_row_count is not None:
            rtxt.add_run("Rows: ").bold = True
            rtxt.add_run(str(trace.final_row_count))
            rtxt.add_run("    ")
        if trace.final_columns:
            rtxt.add_run("Columns: ").bold = True
            rtxt.add_run(", ".join(trace.final_columns))

        doc.add_heading("Preview", level=2)
        _add_preview_table(doc, trace.final_preview, max_rows=12)

        doc.add_heading("Checks", level=2)
        issues = ", ".join(trace.remaining_issues) if trace.remaining_issues else "None"
        pol = (
            "; ".join(trace.remaining_policy_warnings)
            if trace.remaining_policy_warnings
            else "None"
        )
        c = doc.add_paragraph()
        c.add_run("Issues: ").bold = True
        c.add_run(issues)
        c.add_run("\n")
        c.add_run("Policy: ").bold = True
        c.add_run(pol)

    doc.save(str(path))
    return path


# Aliases for callers that still use the old names
write_episode_pdf = write_episode_docx
write_session_pdf = write_session_docx


class EpisodeTraceBuilder:
    """Collects steps; ``build(..., env=)`` turns them into a compact ``EpisodeTrace``."""

    def __init__(self, task: str, mode: str, model_name: Optional[str] = None) -> None:
        self._task = task
        self._mode = mode
        self._model_name = model_name
        self._instruction = ""
        self._snaps: List[_StepSnap] = []
        self._last_obs_after: Optional[DataCleaningObservation] = None

    def set_initial(self, obs: DataCleaningObservation) -> None:
        self._instruction = obs.instruction

    def add_step(
        self,
        obs_before: DataCleaningObservation,
        action: DataCleaningAction,
        obs_after: DataCleaningObservation,
        *,
        parse_ok: bool = True,
        llm_raw: Optional[str] = None,
        heuristic_note: Optional[str] = None,
    ) -> None:
        del obs_before, parse_ok, llm_raw, heuristic_note
        self._snaps.append(_StepSnap(action=action.model_dump(exclude_none=True)))
        self._last_obs_after = obs_after

    def build(
        self,
        terminal_grader_score: Optional[float],
        env: Optional[Any] = None,
    ) -> EpisodeTrace:
        operations_done = [
            f"{i}. {_brief_action_dict(s.action)}"
            for i, s in enumerate(self._snaps, 1)
        ]
        final_row_count: Optional[int] = None
        final_columns: List[str] = []
        final_preview: List[Dict[str, Any]] = []
        if env is not None:
            final_row_count = env.working_row_count()
            final_columns = env.working_column_names()
            final_preview = env.working_preview_records(15)

        remaining_issues: List[str] = []
        remaining_policy_warnings: List[str] = []
        if self._last_obs_after is not None:
            remaining_issues = list(self._last_obs_after.issues)
            remaining_policy_warnings = list(
                getattr(self._last_obs_after, "policy_warnings", []) or []
            )

        return EpisodeTrace(
            task=self._task,
            mode=self._mode,
            model_name=self._model_name,
            instruction=self._instruction,
            terminal_grader_score=terminal_grader_score,
            operations_done=operations_done,
            final_row_count=final_row_count,
            final_columns=final_columns,
            final_preview=final_preview,
            remaining_issues=remaining_issues,
            remaining_policy_warnings=remaining_policy_warnings,
        )
